"""Document Processing Service - PDF text extraction, OCR fallback, DOCX support, and intelligent chunking."""

import io
import re
import uuid
from dataclasses import dataclass

import fitz  # PyMuPDF
import tiktoken
import structlog

from app.config import get_settings

logger = structlog.get_logger()
settings = get_settings()


# Minimum chars per page to consider text extraction successful.
# Below this threshold, we assume the page is scanned/image-based and try OCR.
OCR_THRESHOLD = 30


@dataclass
class Chunk:
    """A chunk of text extracted from a document."""
    chunk_id: str
    chunk_index: int
    text: str
    page_number: int | None
    section_title: str | None
    token_count: int


@dataclass
class ProcessedDocument:
    """Result of processing a document."""
    full_text: str
    chunks: list[Chunk]
    page_count: int
    metadata: dict


class DocumentProcessor:
    """Processes PDF and DOCX documents: extracts text, chunks intelligently."""

    def __init__(self):
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap

    # ── Public API ─────────────────────────────────────────────────────────

    def process_pdf(self, pdf_bytes: bytes, filename: str = "") -> ProcessedDocument:
        """
        Extract text from PDF and chunk it intelligently.

        Pipeline:
        1. Extract text with page mapping (PyMuPDF)
        2. For pages with no/little text, run Tesseract OCR
        3. Detect section headers from layout
        4. Chunk by sections (primary) or sliding window (fallback)
        5. Return structured chunks with metadata
        """
        logger.info("Processing PDF", filename=filename, size_bytes=len(pdf_bytes))

        # Step 1 + 2: Extract text with OCR fallback
        pages, ocr_pages = self._extract_pages(pdf_bytes)
        full_text = "\n\n".join(page["text"] for page in pages)
        page_count = len(pages)

        logger.info(
            "Text extracted",
            pages=page_count,
            chars=len(full_text),
            ocr_pages=ocr_pages,
        )

        if not full_text.strip():
            logger.warning("No text extracted from PDF", filename=filename)
            return ProcessedDocument(
                full_text="",
                chunks=[],
                page_count=page_count,
                metadata={
                    "filename": filename,
                    "page_count": page_count,
                    "chunk_count": 0,
                    "chunking_method": "none",
                    "ocr_pages": ocr_pages,
                    "warning": "No extractable text found",
                },
            )

        # Step 3: Detect sections
        sections = self._detect_sections(pages)

        # Step 4: Chunk
        if sections and len(sections) > 1:
            chunks = self._chunk_by_sections(sections)
            method = "section"
            logger.info("Chunked by sections", chunk_count=len(chunks), section_count=len(sections))
        else:
            chunks = self._chunk_sliding_window(pages)
            method = "sliding_window"
            logger.info("Chunked by sliding window", chunk_count=len(chunks))

        return ProcessedDocument(
            full_text=full_text,
            chunks=chunks,
            page_count=page_count,
            metadata={
                "filename": filename,
                "page_count": page_count,
                "chunk_count": len(chunks),
                "chunking_method": method,
                "ocr_pages": ocr_pages,
            },
        )

    def process_docx(self, docx_bytes: bytes, filename: str = "") -> ProcessedDocument:
        """
        Extract text from a DOCX file and chunk it.

        Uses python-docx to extract paragraphs and tables.
        """
        logger.info("Processing DOCX", filename=filename, size_bytes=len(docx_bytes))

        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(docx_bytes))

            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            full_text = "\n\n".join(paragraphs)

        except ImportError:
            logger.warning("python-docx not installed, falling back to raw decode")
            full_text = docx_bytes.decode("utf-8", errors="replace")
        except Exception as e:
            logger.error("DOCX extraction failed", error=str(e))
            full_text = docx_bytes.decode("utf-8", errors="replace")

        if not full_text.strip():
            return ProcessedDocument(
                full_text="", chunks=[], page_count=1,
                metadata={"filename": filename, "warning": "No text extracted"},
            )

        # Chunk the full text
        chunks = self._split_text(full_text, page_number=1, section_title=None, start_index=0)

        logger.info("DOCX processed", chunks=len(chunks), chars=len(full_text))

        return ProcessedDocument(
            full_text=full_text,
            chunks=chunks,
            page_count=1,
            metadata={
                "filename": filename,
                "page_count": 1,
                "chunk_count": len(chunks),
                "chunking_method": "sliding_window",
            },
        )

    def process_txt(self, text_bytes: bytes, filename: str = "") -> ProcessedDocument:
        """Process a plain text file."""
        text = text_bytes.decode("utf-8", errors="replace").strip()

        if not text:
            return ProcessedDocument(
                full_text="", chunks=[], page_count=1,
                metadata={"filename": filename, "warning": "Empty file"},
            )

        chunks = self._split_text(text, page_number=1, section_title=None, start_index=0)

        return ProcessedDocument(
            full_text=text,
            chunks=chunks,
            page_count=1,
            metadata={
                "filename": filename,
                "page_count": 1,
                "chunk_count": len(chunks),
                "chunking_method": "sliding_window",
            },
        )

    # ── PDF Extraction ─────────────────────────────────────────────────────

    def _extract_pages(self, pdf_bytes: bytes) -> tuple[list[dict], int]:
        """
        Extract text from each page. Falls back to Tesseract OCR for
        pages with little or no text (scanned documents).

        Returns (pages, ocr_page_count).
        """
        pages = []
        ocr_pages = 0

        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")

            for page_num, page in enumerate(doc, 1):
                # First try native text extraction
                text = page.get_text("text").strip()

                if len(text) < OCR_THRESHOLD:
                    # Page likely scanned — try OCR
                    ocr_text = self._ocr_page(page, page_num)
                    if ocr_text and len(ocr_text) > len(text):
                        text = ocr_text
                        ocr_pages += 1

                pages.append({
                    "page_number": page_num,
                    "text": text,
                })

            doc.close()

        except Exception as e:
            logger.error("PDF extraction failed", error=str(e))
            raise

        return pages, ocr_pages

    def _ocr_page(self, page, page_num: int) -> str:
        """
        Run Tesseract OCR on a single PDF page.

        Renders the page to an image, then uses pytesseract to extract text.
        Returns empty string if OCR is unavailable or fails.
        """
        try:
            import pytesseract
            from PIL import Image

            # Render page to image at 300 DPI for good OCR quality
            # Use a zoom factor: 300/72 ≈ 4.17
            zoom = 300 / 72
            matrix = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=matrix, alpha=False)

            # Convert to PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Run Tesseract OCR
            text = pytesseract.image_to_string(img, lang="eng")
            text = text.strip()

            if text:
                logger.info(
                    "OCR extracted text",
                    page=page_num,
                    chars=len(text),
                )

            return text

        except ImportError:
            logger.warning(
                "pytesseract not available, skipping OCR",
                page=page_num,
            )
            return ""

        except Exception as e:
            logger.warning(
                "OCR failed for page",
                page=page_num,
                error=str(e),
            )
            return ""

    # ── Section Detection ──────────────────────────────────────────────────

    def _detect_sections(self, pages: list[dict]) -> list[dict]:
        """
        Detect section headers using common patterns in insurance documents.
        Looks for: all-caps lines, numbered sections, bold-like patterns.
        """
        sections = []
        current_section = {"title": "Document Start", "text": "", "page_number": 1}

        header_patterns = [
            r"^[A-Z][A-Z\s\-]{5,}$",              # ALL CAPS lines (5+ chars)
            r"^(?:SECTION|ARTICLE|PART)\s+\d+",     # SECTION 1, ARTICLE 2, etc.
            r"^\d+\.\s+[A-Z]",                       # 1. Title format
            r"^[IVXLC]+\.\s+",                       # Roman numeral sections
            r"^(?:COVERAGE|EXCLUSION|CONDITION|DEFINITION|ENDORSEMENT)",
        ]

        for page in pages:
            lines = page["text"].split("\n")
            for line in lines:
                line_stripped = line.strip()
                if not line_stripped:
                    continue

                is_header = any(re.match(pattern, line_stripped) for pattern in header_patterns)

                if is_header and len(line_stripped) < 100:
                    if current_section["text"].strip():
                        sections.append(current_section)
                    current_section = {
                        "title": line_stripped,
                        "text": "",
                        "page_number": page["page_number"],
                    }
                else:
                    current_section["text"] += line + "\n"

        if current_section["text"].strip():
            sections.append(current_section)

        return sections

    # ── Chunking ───────────────────────────────────────────────────────────

    def _chunk_by_sections(self, sections: list[dict]) -> list[Chunk]:
        """Chunk by detected sections. Split large sections with sliding window."""
        chunks = []
        chunk_index = 0

        for section in sections:
            text = section["text"].strip()
            token_count = len(self.tokenizer.encode(text))

            if token_count <= self.chunk_size:
                chunks.append(Chunk(
                    chunk_id=str(uuid.uuid4()),
                    chunk_index=chunk_index,
                    text=text,
                    page_number=section["page_number"],
                    section_title=section["title"],
                    token_count=token_count,
                ))
                chunk_index += 1
            else:
                sub_chunks = self._split_text(
                    text,
                    page_number=section["page_number"],
                    section_title=section["title"],
                    start_index=chunk_index,
                )
                chunks.extend(sub_chunks)
                chunk_index += len(sub_chunks)

        return chunks

    def _chunk_sliding_window(self, pages: list[dict]) -> list[Chunk]:
        """Fallback: chunk using sliding window across all pages."""
        chunks = []
        chunk_index = 0

        for page in pages:
            text = page["text"].strip()
            if not text:
                continue

            sub_chunks = self._split_text(
                text,
                page_number=page["page_number"],
                section_title=None,
                start_index=chunk_index,
            )
            chunks.extend(sub_chunks)
            chunk_index += len(sub_chunks)

        return chunks

    def _split_text(
        self,
        text: str,
        page_number: int | None,
        section_title: str | None,
        start_index: int,
    ) -> list[Chunk]:
        """Split text into overlapping chunks by token count."""
        tokens = self.tokenizer.encode(text)
        chunks = []
        start = 0

        while start < len(tokens):
            end = min(start + self.chunk_size, len(tokens))
            chunk_tokens = tokens[start:end]
            chunk_text = self.tokenizer.decode(chunk_tokens)

            chunks.append(Chunk(
                chunk_id=str(uuid.uuid4()),
                chunk_index=start_index + len(chunks),
                text=chunk_text.strip(),
                page_number=page_number,
                section_title=section_title,
                token_count=len(chunk_tokens),
            ))

            if end >= len(tokens):
                break
            start += self.chunk_size - self.chunk_overlap

        return chunks

    def count_tokens(self, text: str) -> int:
        """Count tokens in a text string."""
        return len(self.tokenizer.encode(text))